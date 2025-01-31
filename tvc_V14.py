import argparse
import pandas as pd
from docx import Document
from docx.shared import Cm, Pt, Inches
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION
from utils import *
from json_access import *
import json
import os

def ecrire_fichier_json(feuille_source_resultat, tourActif, top_poules, poules):
    """
    Écrit les données des poules dans un fichier JSON avec une structure détaillée.
    """
    # Extraire le numéro du tour depuis la chaîne tourActif (e.g., "tour3" -> 3)
    round_number = int(tourActif.replace("tour", ""))

    data = {
        "category": feuille_source_resultat,
        "round": round_number,
        "poules": {}
    }

    # Ajouter les données des tops
    #print("top_poules:", top_poules)
    #for titre, joueurs in top_poules.items():
    for idx, poule in enumerate(top_poules, start=1):
        for joueur in poule:
            print("JOUEUR:", joueur)
            
    for idx, poule in enumerate(top_poules, start=1):
        data["poules"][f"Top {idx}"] = {
            "joueurs": [
                {
                    "nom": joueur["nom"],
                    "prenom": joueur["prenom"],
                    "club": joueur["club"],
                    "classement": joueur["classement"],
                    "points": joueur["points"],
                }
                for joueur in poule
            ]
        }

    # Ajouter les données des poules classiques
    for idx, poule in enumerate(poules, start=1):
        data["poules"][f"Poule {idx}"] = {
            "joueurs": [
                {
                    "nom": joueur["nom"],
                    "prenom": joueur["prenom"],
                    "club": joueur["club"],
                    "classement": joueur["classement"],
                    "points": joueur["points"],
                }
                for joueur in poule
            ]
        }

    # Nom du fichier JSON basé sur la feuille et le tour
    fichier_json = f"{feuille_source_resultat}_{tourActif}.json"
    
    save_poules(fichier_json, data)

def charger_joueurs(fichier_entree, feuille_source_resultat):
    # Lire le fichier Excel en spécifiant le nom de la feuille et en ignorant les premières lignes
    joueurs = pd.read_excel(fichier_entree, sheet_name=feuille_source_resultat, header=6)

    # Renommer les colonnes pour correspondre aux attentes du script
    mapping_colonnes = {
        "NOM": "nom",
        "PRENOM": "prenom",
        "Sexe": "sexe",
        "Naiss": "naissance",
        "Club": "club",
        "Class.": "classement",
        "Pr. T1": "Tour1",
        "Points T1": "PtsT1",
        "Pr. T2": "Tour2",
        "Points T2": "PtsT2",
        "Pr. T3": "Tour3",
        "Points T3": "PtsT3",
        "Pr. T4": "Tour4",
        "Points T4": "PtsT4",
        "TOTAL": "points"
    }
    joueurs.rename(columns=mapping_colonnes, inplace=True)

    # Remplacer les colonnes manquantes par des colonnes vides si nécessaire
    for col in ["Tour1", "Tour2", "Tour3", "Tour4"]:
        if col not in joueurs.columns or joueurs[col].isnull().all():
            joueurs[col] = ""

    return joueurs    

def filtrer_et_trier_joueurs(joueurs):
    mapping = [
        ("Tour4", "Nb Joueurs Inscrits Tour4: ",  "points",     True),
        ("Tour3", "Nb Joueurs Inscrits Tour3: ",  "points",     True),
        ("Tour2", "Nb Joueurs Inscrits Tour2: ",  "points",     True),
        ("Tour1", "Nb Joueurs Inscrits Tour 1: ", "classement", False),
    ]
    
    # 1) Préparer les colonnes Tour1..Tour4
    for col in ["Tour1", "Tour2", "Tour3", "Tour4"]:
        joueurs[col] = joueurs[col].fillna("").astype(str)
    
    # 2) On cherche le premier tour qui contient un 'X'
    for col, msg, sort_col, active_bool in mapping:
        if joueurs[col].str.contains('X', case=False, na=False).any():
            # On filtre le DataFrame sur la colonne correspondante
            filtered = joueurs[joueurs[col].str.contains('X', case=False, na=False)]
            
            nb_joueurs = len(filtered)
            filtered = filtered.sort_values(by=sort_col, ascending=False)
            return filtered, active_bool, col.lower(), nb_joueurs
        
    return joueurs, False, None, 0
    
def construire_poules_restantes(autres_joueurs, nb_poules_3, nb_poules_4):
    poules = [[] for _ in range(nb_poules_4 + nb_poules_3)]
    index_poule = 0
    direction = 1  # 1 pour avancer, -1 pour reculer
    step_count = 0

    for joueur in autres_joueurs.to_dict('records'):
        poules[index_poule].append(joueur)

        if direction == 1:
            if index_poule == len(poules) - 1:
                if (step_count == 0):
                    step_count = 1
                else:
                    direction = -1
                    index_poule -= 1
                    step_count = 0
            else:
                index_poule += 1
        else:
            if (index_poule == 0):
                if (step_count == 0):
                    step_count = 1
                else:
                    direction = 1
                    index_poule += 1
                    step_count = 0                    
            else:
                index_poule -= 1
    return poules

def club_present_dans_poule(poules, index_poule, club, nom="", prenom=""):
    found = False
    for joueur in poules[index_poule]:
        if (joueur['club'] == club and (joueur['nom'] != nom or joueur['prenom'] != prenom)):
            found = True
    return found

def deplacer_index(index_poule, nb_poules, direction, step_count, position):
        if direction == 1:
            if index_poule == nb_poules - 1:
                position = position + 1
                direction = -1
            else:
                index_poule += 1
        else:
            if index_poule == 0:
                position = position + 1
                direction = 1
            else:
                index_poule -= 1
        return index_poule, direction, step_count, position
         
def position_vide(poule, position):
    if (len(poule) < (position+1)):
        return True
    else:
        return False

def chercher_si_echange_possible(poules, club, direction, index_courant, position, nb_poules):
    idx = index_courant
    pos = position
    fini = False
    liste_poules = []
    
    while ( (idx >= 0 and direction ==- 1) or (idx < nb_poules and direction == 1) ):
        if (not club_present_dans_poule(poules, idx, club)):
            print("club potentiellement non présent, verif que le club du joueur interverti pas présent dans poule courante")
            joueur = poules[idx][position]
            club_present = club_present_dans_poule(poules, index_courant, joueur['club'])
            if (not club_present):
                print ("possibilité d'intervertir avec ", joueur['nom'], " ", joueur['prenom'], " ", joueur['club'])
                return idx
            #liste_poules.append(idx)
        if (direction == 1):
            idx = idx + 1
        else:
            idx = idx - 1
    return -1
    
def echanger_joueurs(poules, index_poule1, position1, index_poule2, position2):
    print("Echange ", poules[index_poule1][position1], " with ", poules[index_poule2])
    poules[index_poule1][position1], poules[index_poule2][position2] = \
        poules[index_poule2][position2], poules[index_poule1][position1]
    
def chercher_poule(poules, club, direction, index_courant, position, nb_poules):
    idx = index_courant
    pos = position
    fini = False
    liste_poules = []
    
    while ( (idx >= 0 and direction ==- 1) or (idx < nb_poules and direction == 1) ):
        if (position_vide(poules[idx], position) and not club_present_dans_poule(poules, idx, club)):
            liste_poules.append(idx)
        if (direction == 1):
            idx = idx + 1
        else:
            idx = idx - 1
    return liste_poules
    
    
'''def construire_tops(joueurs, nb_poules_top):
    joueurs_top = joueurs.iloc[:nb_poules_top * 4]
    autres_joueurs = joueurs.iloc[nb_poules_top * 4:].reset_index(drop=True)
    top_poules = {f"Top {i + 1}": [] for i in range(nb_poules_top)}
    ordre = [f"Top {i % nb_poules_top + 1}" for i in range(nb_poules_top * 4)]
    for i, joueur in enumerate(joueurs_top.itertuples(index=False)):
        top_poules[ordre[i]].append(joueur)
    return top_poules, autres_joueurs   
'''

def construire_tops(joueurs, nb_poules_top):
    index_poule = 0
    direction = 1  # 1 pour avancer, -1 pour reculer
    step_count = 0
    position = 0
    
    joueurs_top = joueurs.iloc[:nb_poules_top * 4]
    autres_joueurs = joueurs.iloc[nb_poules_top * 4:].reset_index(drop=True)

    top_poules = [[] for _ in range(nb_poules_top)]
    for joueur in joueurs_top.to_dict('records'):
        if (position == 0):
            top_poules[index_poule].append(joueur)
            index_poule, direction, step_count, position = deplacer_index(index_poule, nb_poules_top, direction, step_count, position)
        else:
            while (not position_vide(top_poules[index_poule], position)):
                index_poule, direction, step_count, position = deplacer_index(index_poule, nb_poules_top, direction, step_count, position)
            if club_present_dans_poule(top_poules, index_poule, joueur['club'], joueur['nom'], joueur['prenom']):
                liste_index = chercher_poule(top_poules, joueur['club'], direction, index_poule, position, nb_poules_top)
                if (len(liste_index) > 0):
                    top_poules[liste_index[0]].append(joueur)
                    print("joueur ", joueur['nom'], " ", joueur['prenom'], " ",  joueur['club'], " ajouté dans poule ", liste_index[0])
                else:
                    idx = chercher_si_echange_possible(top_poules, joueur['club'], -direction, index_poule, position, nb_poules_top)
                    top_poules[index_poule].append(joueur)
                    if (idx == -1):
                        print("joueur ", joueur['nom'], " ", joueur['prenom'], " ",  joueur['club'], " ajouté dans poule ", index_poule)
                        index_poule, direction, step_count, position = deplacer_index(index_poule, nb_poules_top, direction, step_count, position)
                    else:
                        echanger_joueurs(top_poules, index_poule, position, idx, position)
            else:
                top_poules[index_poule].append(joueur)
                print("joueur ", joueur['nom'], " ", joueur['prenom'], " ",  joueur['club'], " ajouté dans poule ", index_poule)
                index_poule, direction, step_count, position = deplacer_index(index_poule, nb_poules_top, direction, step_count, position)
                 
    return top_poules, autres_joueurs
    
def placer_joueurs(autres_joueurs, nb_poules_3, nb_poules_4):
    index_poule = 0
    direction = 1  # 1 pour avancer, -1 pour reculer
    step_count = 0
    position =0
    poules = [[] for _ in range(nb_poules_4 + nb_poules_3)]
    
    nb_poules = nb_poules_3 + nb_poules_4
    for joueur in autres_joueurs.to_dict('records'):
        if (position == 0):
            print("joueur ", joueur['nom'], " ", joueur['prenom'], " ",  joueur['club'], " ajouté dans poule ", index_poule)
            poules[index_poule].append(joueur)
            index_poule, direction, step_count, position = deplacer_index(index_poule, nb_poules, direction, step_count, position)    
            
        else:
            while (not position_vide(poules[index_poule], position)):
                index_poule, direction, step_count, position = deplacer_index(index_poule, nb_poules, direction, step_count, position)
            if club_present_dans_poule(poules, index_poule, joueur['club'], joueur['nom'], joueur['prenom']):
                liste_index = chercher_poule(poules, joueur['club'], direction, index_poule, position, nb_poules)
                if (len(liste_index) > 0):
                    poules[liste_index[0]].append(joueur)
                    print("joueur ", joueur['nom'], " ", joueur['prenom'], " ",  joueur['club'], " ajouté dans poule ", liste_index[0])
                else:
                    idx = chercher_si_echange_possible(poules, joueur['club'], -direction, index_poule, position, nb_poules)
                    poules[index_poule].append(joueur)
                    if (idx == -1):
                        print("joueur ", joueur['nom'], " ", joueur['prenom'], " ",  joueur['club'], " ajouté dans poule ", index_poule)
                        index_poule, direction, step_count, position = deplacer_index(index_poule, nb_poules, direction, step_count, position)
                    else:
                        echanger_joueurs(poules, index_poule, position, idx, position)
            else:
                poules[index_poule].append(joueur)
                print("joueur ", joueur['nom'], " ", joueur['prenom'], " ",  joueur['club'], " ajouté dans poule ", index_poule)
                index_poule, direction, step_count, position = deplacer_index(index_poule, nb_poules, direction, step_count, position)
                  
    return poules

def ajuster_espace_autour_tableau(table):
    # Paragraphe avant le tableau
    prev_paragraph = table._element.getprevious()
    if prev_paragraph is not None and prev_paragraph.tag.endswith("p"):
        prev_p = table._parent.paragraphs[-1]
        prev_p.paragraph_format.space_after = Pt(0)
        prev_p.paragraph_format.space_before = Pt(0)

    # Paragraphe après le tableau
    next_paragraph = table._element.getnext()
    if next_paragraph is not None and next_paragraph.tag.endswith("p"):
        next_p = table._parent.paragraphs[-1]
        next_p.paragraph_format.space_after = Pt(0)
        next_p.paragraph_format.space_before = Pt(0)

            
def ecrire_poule_on_frame(doc, titre, joueurs, is_top=False):
    oppositions = [(1, 4), (2, 3), (1, 3), (2, 4), (1, 2), (2, 4)]
    heading = titre if is_top else f"Poule {titre}"

    paragraph = doc.add_paragraph()
    paragraph.text = f'{titre}'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Style du titre
    run = paragraph.runs[0]
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.bold = True

    #table = doc.add_table(rows=1, cols=5)
    table = doc.add_table(rows=1, cols=4)
    ajuster_espace_autour_tableau(table)
    table.autofit = False
    table.style = 'Table Grid'

    table.columns[0].width = Cm(0.5)
    table.columns[1].width = Cm(4)
    table.columns[2].width = Cm(3.5)
    table.columns[3].width = Cm(1.5)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = ' '
    hdr_cells[1].text = 'Nom'
    hdr_cells[2].text = 'Prenom'
    hdr_cells[3].text = 'Club'
    #hdr_cells[4].text = 'Points'

    for idx, joueur in enumerate(joueurs, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = joueur['nom'] if isinstance(joueur, dict) else joueur.nom
        row_cells[2].text = joueur['prenom'] if isinstance(joueur, dict) else joueur.prenom
        row_cells[3].text = joueur['club'] if isinstance(joueur, dict) else joueur.club
        #row_cells[4].text = str(joueur['points']) if isinstance(joueur, dict) else str(joueur.points)

    table2 = doc.add_table(rows=1, cols=6 + len(joueurs))
    table2.style = 'Table Grid'
    ajuster_espace_autour_tableau(table2)
    table2.columns[0].width = Cm(1)
   
    for i in range(1, 6 + len(joueurs)):
            if i < len(table2.columns):
                table2.columns[i].width = Cm(0.8)

    hdr_cells = table2.rows[0].cells
    hdr_cells[6].text = "1"
    hdr_cells[7].text = "2"
    hdr_cells[8].text = "3"
    if len(joueurs) == 4:
            hdr_cells[9].text = "4"

    for z in range(len(oppositions)):
            x, y = oppositions[z]
            if x <= len(joueurs) and y <= len(joueurs):
                row_cells = table2.add_row().cells
                rencontre = f'{x} vs {y}'
                row_cells[0].text = rencontre
            for k in range (1,len(joueurs) + 1):
                if (k != x and k != y and y<=len(joueurs)):
                    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
                    row_cells[k+5]._tc.get_or_add_tcPr().append(shading_elm_1)
                
    
    table3 = doc.add_table(rows=1, cols=2)
    ajuster_espace_autour_tableau(table3)
    table3.style = 'Table Grid'
    table3.columns[0].width = Cm(1)
    row_cells = table3.rows[0].cells
    for index in range(1, len(joueurs) + 1):
            if (index == 1):
                label = "1er"
            else:
                label = f"{index}eme"
            row_cells[0].text = f'{label}'
            if index < len(joueurs):
                row_cells = table3.add_row().cells

def create_word_file_with_frames(doc, groupes, categorie, first_poule, is_top, position=0, table=None):
    # Configurer les marges de la page
    section = doc.sections[0]
    section.page_width = Inches(8.27)  # Largeur de la page (A4)
    section.page_height = Inches(11.69)  # Hauteur de la page (A4)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    # Si aucune table n'existe, en créer une nouvelle
    if table is None:
        table = doc.add_table(rows=3, cols=2)
        table.autofit = False  # Désactiver l'ajustement automatique
        for column in table.columns:
            column.width = Inches(3.75)

    for i, joueurs in enumerate(groupes, start=1):
        # Ajouter une nouvelle page après avoir rempli la table
        if position == 0 and i > 0:
            if not first_poule:
                doc.add_page_break()
                table = doc.add_table(rows=3, cols=2)
                table.autofit = False
                for column in table.columns:
                    column.width = Inches(3.75)
            else:
                first_poule = False

        # Déterminer la cellule où écrire
        if position == 0:  # Haut gauche
            cell = table.cell(0, 0)
        elif position == 1:  # Haut droite
            cell = table.cell(0, 1)
        elif position == 2:  # Milieu gauche
            cell = table.cell(1, 0)
        elif position == 3:  # Milieu droite
            cell = table.cell(1, 1)
        elif position == 4:  # Bas gauche
            cell = table.cell(2, 0)
        elif position == 5:  # Bas droite
            cell = table.cell(2, 1)
       
        if is_top:
            titre=f"Top {i}"
        else:
            titre=f"Poule {i}"
        titrePoule = f"{titre} {categorie}"
        ecrire_poule_on_frame(cell, titrePoule, joueurs, is_top)

        # Mettre à jour la position
        position = (position + 1) % 6  # Cycle entre 0 et 5

    return position, first_poule, table  # Retourner la position actuelle et la table en cours

def ecrire_document_word(fichier_sortie, feuille_source_resultat, top_poules, poules):
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = "Arial"
    font.size = Pt(7)

    first_poule = True
    
    # Initialiser la position à 0 pour démarrer avec les poules de top
    position, first_poule, table = create_word_file_with_frames(doc, top_poules, feuille_source_resultat, first_poule, is_top=True)

    # Écrire les poules normales à partir de la dernière position et table
    position, _, _ = create_word_file_with_frames(doc, poules, feuille_source_resultat, first_poule, is_top=False, position=position, table=table)

    fichier_sortie = lire_word_name(fichier_sortie)
    doc.save(fichier_sortie)
    print(f'Fichier Word généré: {fichier_sortie}')

def repartir_poules_sans_conflit(fichier_entree, feuille_source_resultat, nb_poules_top=None, nb_poules_3=None, nb_poules_4=None):
    joueurs = charger_joueurs(fichier_entree, feuille_source_resultat)
    joueurs, has_top, tourActif, nbJoueurs = filtrer_et_trier_joueurs(joueurs)


    if has_top and nb_poules_top is None:
        nb_poules_top = 2
    elif (nb_poules_top is None):
        nb_poules_top = 0

    if (nb_poules_3 != None and nb_poules_4 != None):
        nbTotal = nb_poules_top*4 + nb_poules_3*3 + nb_poules_4*4
        if (nbTotal != nbJoueurs):
            print(nbJoueurs, " joueurs inscrits, erreur dans la definition des parametres (nb_poules_top: ", nb_poules_top, ", nb_poules_3: ", nb_poules_3, ", nb_poules_4", nb_poules_4)
            exit()
        
    if has_top:
        top_poules, autres_joueurs = construire_tops(joueurs, nb_poules_top)
    else:
        top_poules = {}
        autres_joueurs = joueurs

    if nb_poules_3 is None or nb_poules_4 is None:
        total_joueurs = int(len(autres_joueurs))
        nb_poules_4 = total_joueurs // 4
        nb_poules_3 = (total_joueurs % 4) // 3
        
        while ( (nb_poules_4*4 + nb_poules_3*3) < total_joueurs):
            nb_poules_4 = nb_poules_4 - 1
            nb_poules_3 = int( (total_joueurs - (nb_poules_4 *4))/3)
        
    print ("Nb Poules de 3", nb_poules_3, ", Nb Poules de 4", nb_poules_4)

    poules = placer_joueurs(autres_joueurs, nb_poules_3, nb_poules_4)
    fichierWord=f'{feuille_source_resultat}_{tourActif}.docx'
    ecrire_document_word(fichierWord, feuille_source_resultat, top_poules, poules)
    # Générer le fichier JSON
    ecrire_fichier_json(feuille_source_resultat, tourActif, top_poules, poules)

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Repartir les poules en fonction des joueurs et des options.")
    parser.add_argument("Categorie", type=str, help="Nom de la feuille source dans le fichier Excel.")
    parser.add_argument("--nb_poules_top", type=int, default=None, help="Nombre de poules Top. Optionnel.")
    parser.add_argument("--nb_poules_3", type=int, default=None, help="Nombre de poules de 3 joueurs. Optionnel.")
    parser.add_argument("--nb_poules_4", type=int, default=None, help="Nombre de poules de 4 joueurs. Optionnel.")
    args = parser.parse_args()
    
    fichier_entree = lire_excel_name()
    print(f"Fichier complet : {fichier_entree}")
        
    nb_tops_defaut = int(charger_config_categorie("config.ini", args.Categorie))
        
    number_of_tops=args.nb_poules_top
    if args.nb_poules_top == None:
        number_of_tops = nb_tops_defaut
   
    print(f"number_of_tops: {number_of_tops}")
   
    repartir_poules_sans_conflit(
        fichier_entree,
        feuille_source_resultat=args.Categorie,
        nb_poules_top=number_of_tops,
        nb_poules_3=args.nb_poules_3,
        nb_poules_4=args.nb_poules_4)
