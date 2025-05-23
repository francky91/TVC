import argparse
import pandas as pd
from docx import Document
from docx.shared import Cm, Pt, Inches
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn 
from utils import *
from json_access import *
import json
import os

def ecrire_fichier_json(feuille_source_resultat, tourActif, top_poules, poules, save_to_file=True):
    """
    Écrit les données des poules dans un fichier JSON et retourne le dictionnaire JSON pour tests.
    """
    round_number = int(tourActif.replace("tour", ""))

    data = {
        "category": feuille_source_resultat,
        "round": round_number,
        "poules": {}
    }

    # Ajouter les données des tops
    for idx, poule in enumerate(top_poules, start=1):
        data["poules"][f"Top {idx}"] = {
            "joueurs": [
                {
                    "nom": joueur["nom"],
                    "prenom": joueur["prenom"],
                    "dossard": str(int(joueur["dossard"])),
                    "club": joueur["club"],
                    "classement": joueur["classement"],
                    "points": joueur["points"],
                }
                for joueur in poule
            ]
        }
        print("")

    # Ajouter les données des poules classiques
    for idx, poule in enumerate(poules, start=1):
        data["poules"][f"Poule {idx}"] = {
            "joueurs": [
                {
                    "nom": joueur["nom"],
                    "prenom": joueur["prenom"],
                    "club": joueur["club"],
                    "dossard": str(int(joueur["dossard"])),
                    "classement": joueur["classement"],
                    "points": joueur["points"],
                }
                for joueur in poule
            ]
        }

    if save_to_file:
        fichier_json = f"{feuille_source_resultat}_{tourActif}.json"
        save_poules(fichier_json, data)

    return data  # Retourne le JSON pour faciliter les tests


def charger_joueurs(fichier_entree, feuille_source_resultat):
    # Lire le fichier Excel en spécifiant le nom de la feuille et en ignorant les premières lignes
    joueurs = pd.read_excel(fichier_entree, sheet_name=feuille_source_resultat, header=6)

    # Renommer les colonnes pour correspondre aux attentes du script
    mapping_colonnes = {
        "Dossard": "dossard",
        "NOM": "nom",
        "PRENOM": "prenom",
        "Sexe": "sexe",
        "Naiss": "naissance",
        "Club": "club",
        "Class.": "classement",
        "Pr. T1": "Tour1",
        "Points T1": "PtsT1",
        "Pr. T2": "Tour2",
        "Points T2": "PtsT2",
        "Pr. T3": "Tour3",
        "Points T3": "PtsT3",
        "Pr. T4": "Tour4",
        "Points T4": "PtsT4",
        "TOTAL": "points"
    }
    joueurs.rename(columns=mapping_colonnes, inplace=True)

    # Remplacer les colonnes manquantes par des colonnes vides si nécessaire
    for col in ["Tour1", "Tour2", "Tour3", "Tour4"]:
        if col not in joueurs.columns or joueurs[col].isnull().all():
            joueurs[col] = ""

    return joueurs    

'''def filtrer_et_trier_joueurs(joueurs, nb_poules_tops):
    # mapping = (colonne de participation, texte, critère principal, booléen, liste des colonnes de points par tour)
    mapping = [
        ("Tour4", "Nb Joueurs Inscrits Tour4: ",  "points",   ["PtsT1", "PtsT2", "PtsT3"]),
        ("Tour3", "Nb Joueurs Inscrits Tour3: ",  "points",  ["PtsT1", "PtsT2"]),
        ("Tour2", "Nb Joueurs Inscrits Tour2: ",  "points",  ["PtsT1"]),
        ("Tour1", "Nb Joueurs Inscrits Tour 1: ", "classement", []),
    ]
    
    # 1) Préparer (remplir NaN, forcer en str) les colonnes Tour1..Tour4 si besoin
    for col in ["Tour1", "Tour2", "Tour3", "Tour4"]:
        joueurs[col] = joueurs[col].fillna("").astype(str)
    
    # 2) On cherche la première colonne 'TourX' où il y a un 'X'
    for col, msg, sort_col, liste_col_points in mapping:
        if joueurs[col].str.contains('X', case=False, na=False).any():
            # On ne garde que les joueurs qui ont un 'X' dans cette colonne
            filtered = joueurs[joueurs[col].str.contains('X', case=False, na=False)].copy()
            
            nb_joueurs = len(filtered)
            
            # Si le critère principal est 'points', on veut un tri à deux niveaux :
            #   1) Descendant sur 'points'
            #   2) Descendant sur le maximum de points marqués dans les tours déjà joués
            if nb_poules_tops == 0:
                sort_col = "classement"
            if sort_col == "points" and len(liste_col_points) > 0:
                # Créer une colonne 'max_points' = maximum sur les tours déjà joués
                filtered['max_points'] = filtered[liste_col_points].max(axis=1)
                
                # Tri multi-critères : d'abord 'points' desc, puis 'max_points' desc
                filtered = filtered.sort_values(
                    by=['points', 'max_points'], 
                    ascending=[False, False]
                )
            else:
                # Sinon, on fait le tri simple
                print("sort_col: ", sort_col)
                filtered = filtered.sort_values(by=sort_col, ascending=False)

            return filtered, col.lower(), nb_joueurs
        
    # Si aucun 'X' n'a été trouvé dans Tour1..Tour4, on renvoie le DataFrame original sans filtre
    return joueurs, None, 0
'''

def filtrer_et_trier_joueurs(joueurs, nb_poules_tops):
    # mapping = (colonne de participation, libellé, critère principal, liste des colonnes de points à utiliser pour le tri secondaire)
    mapping = [
        ("Tour4", "Nb Joueurs Inscrits Tour4: ", "points", ["PtsT3"]),  # Utilise uniquement le score du Tour3
        ("Tour3", "Nb Joueurs Inscrits Tour3: ", "points", ["PtsT2"]),  # Utilise uniquement le score du Tour2
        ("Tour2", "Nb Joueurs Inscrits Tour2: ", "points", ["PtsT1"]),  # Utilise uniquement le score du Tour1
        ("Tour1", "Nb Joueurs Inscrits Tour 1: ", "classement", [])
    ]
    
    # 1) Remplacer les NaN par une chaîne vide
    for col in ["Tour1", "Tour2", "Tour3", "Tour4"]:
        joueurs[col] = joueurs[col].fillna("").astype(str)
    
    # 2) Parcourir le mapping pour trouver le premier TourX où il y a un 'X'
    for col, msg, sort_col, liste_col_points in mapping:
        if joueurs[col].str.contains('X', case=False, na=False).any():
            # Garder uniquement les joueurs présents dans ce tour
            filtered = joueurs[joueurs[col].str.contains('X', case=False, na=False)].copy()
            nb_joueurs = len(filtered)

            # Définir le seuil : par exemple, pour 1 top, on garde les 4 premiers ; pour 2 tops, 8 premiers.
            threshold = nb_poules_tops * 4

            # Si nb_poules_tops == 0, on trie par "classement"
            if nb_poules_tops == 0:
                sort_col = "classement"

            # --- CAS : tri par points (desc) ---
            if sort_col == "points" and len(liste_col_points) > 0:
                # Calculer le score du tour précédent
                filtered["prev_tour"] = filtered[liste_col_points].max(axis=1)
                
                # Étape A: tri par points décroissant uniquement
                stepA = filtered.sort_values(by="points", ascending=False)
                if len(stepA) >= threshold + 1:
                    if stepA.iloc[threshold - 1]["points"] > stepA.iloc[threshold]["points"]:
                        return stepA, col.lower(), nb_joueurs

                # Étape B: ajouter le critère du tour précédent
                print ("In Step B")
                stepB = filtered.sort_values(by=["points", "prev_tour"], ascending=[False, False])
                if len(stepB) >= threshold + 1:
                    # Si le joueur à la position seuil est clairement supérieur à celui qui suit
                    if (stepB.iloc[threshold - 1]["points"] > stepB.iloc[threshold]["points"]) or \
                       (stepB.iloc[threshold - 1]["points"] == stepB.iloc[threshold]["points"] and
                        stepB.iloc[threshold - 1]["prev_tour"] > stepB.iloc[threshold]["prev_tour"]):
                        return stepB, col.lower(), nb_joueurs

                # Étape C: ajouter la date de naissance (du plus jeune au plus âgé)
                # Ici, pour obtenir le plus jeune en premier, on trie par date décroissante
                print ("In Step C, columns", filtered.columns)
                if "naissance" in filtered.columns:
                    print ("In Step C-2")
                    stepC = filtered.sort_values(
                        by=["points", "prev_tour", "naissance"],
                        ascending=[False, False, False]
                    )
                    return stepC, col.lower(), nb_joueurs
                else:
                    return stepB, col.lower(), nb_joueurs
            else:
                # Cas de Tour1 ou autre tri simple
                filtered_simple = filtered.sort_values(by=sort_col, ascending=False)
                return filtered_simple, col.lower(), nb_joueurs

    # Si aucun 'X' n'est trouvé, retourner le DataFrame original
    return joueurs, None, 0
    
def construire_poules_restantes(autres_joueurs, nb_poules_3, nb_poules_4):
    poules = [[] for _ in range(nb_poules_4 + nb_poules_3)]
    index_poule = 0
    direction = 1  # 1 pour avancer, -1 pour reculer
    step_count = 0

    for joueur in autres_joueurs.to_dict('records'):
        poules[index_poule].append(joueur)

        if direction == 1:
            if index_poule == len(poules) - 1:
                if (step_count == 0):
                    step_count = 1
                else:
                    direction = -1
                    index_poule -= 1
                    step_count = 0
            else:
                index_poule += 1
        else:
            if (index_poule == 0):
                if (step_count == 0):
                    step_count = 1
                else:
                    direction = 1
                    index_poule += 1
                    step_count = 0                    
            else:
                index_poule -= 1
    return poules

def club_present_dans_poule(poules, index_poule, club, nom="", prenom=""):
    found = False
    for joueur in poules[index_poule]:
        if (joueur['club'] == club and (joueur['nom'] != nom or joueur['prenom'] != prenom)):
            found = True
    return found

def deplacer_index(index_poule, nb_poules, direction, step_count, position):
        if direction == 1:
            if index_poule == nb_poules - 1:
                position = position + 1
                direction = -1
            else:
                index_poule += 1
        else:
            if index_poule == 0:
                position = position + 1
                direction = 1
            else:
                index_poule -= 1
        return index_poule, direction, step_count, position
         
def position_vide(poule, position):
    if (len(poule) < (position+1)):
        return True
    else:
        return False


    
def echanger_joueurs(poules, index_poule1, position1, index_poule2, position2):
    print("Echange poules[index_poule1][position1] with poules[index_poule2][position2]")
    print("poules[index_poule1][position1]:", poules[index_poule1][position1], ",poules[index_poule2][position2]:",poules[index_poule2][position2])
    poules[index_poule1][position1], poules[index_poule2][position2] = \
        poules[index_poule2][position2], poules[index_poule1][position1]
    
def chercher_poule(poules, club, direction, index_courant, position, nb_poules, taille_poules):
    idx = index_courant
    pos = position
    fini = False
    liste_poules = []
    
    while ( (idx >= 0 and direction ==- 1) or (idx < nb_poules and direction == 1) ):
        #if (position_vide(poules[idx], position) and not club_present_dans_poule(poules, idx, club)):
        x = len(poules[idx])
        print(f"poule {idx}, taille poule: {taille_poules[idx]}, len: {x}")
        if (len(poules[idx]) < taille_poules[idx] and not club_present_dans_poule(poules, idx, club)):
            print(f"append {idx}")
            liste_poules.append(idx)
        if (direction == 1):
            idx = idx + 1
        else:
            idx = idx - 1
    return liste_poules

def chercher_si_echange_possible(poules, club, direction, index_courant, position, nb_poules, taille_poules):
    idx = index_courant
    pos = position
    fini = False
    liste_poules = []
    
    while ( (idx >= 0 and direction ==- 1) or (idx < nb_poules and direction == 1) ):
        #if (not club_present_dans_poule(poules, idx, club)):
        if (position < taille_poules[idx] and not club_present_dans_poule(poules, idx, club)):
            print("club potentiellement non présent, verif que le club du joueur interverti pas présent dans poule courante")
            joueur = poules[idx][position]
            club_present = club_present_dans_poule(poules, index_courant, joueur['club'])
            if (not club_present):
                print ("possibilité d'intervertir avec ", joueur['nom'], " ", joueur['prenom'], " ", joueur['club'])
                return idx
            #liste_poules.append(idx)
        if (direction == 1):
            idx = idx + 1
        else:
            idx = idx - 1
    return -1

def construire_tops(joueurs, nb_poules_top):
    index_poule = 0
    direction = 1  # 1 pour avancer, -1 pour reculer
    step_count = 0
    position = 0
    
    if nb_poules_top == 0:
        return [], joueurs.copy()
        
    joueurs_top = joueurs.iloc[:nb_poules_top * 4]
    autres_joueurs = joueurs.iloc[nb_poules_top * 4:].reset_index(drop=True)
    top_poules = [[] for _ in range(nb_poules_top)]
    
    print(f"joueurs_top: {joueurs_top}")
    if (nb_poules_top == 2):
        poule_top1, poule_top2 = construire_tops_2_poules(joueurs_top)
        top_poules[0] = poule_top1
        top_poules[1] = poule_top2
    else:
        top_poules[0] = joueurs_top.to_dict('records')
    return top_poules, autres_joueurs



def construire_club_counts(pouleA, pouleB):
    """
    Retourne un dict : { "ClubX": [countEnA, countEnB], ... }
    """
    from collections import defaultdict
    counts = defaultdict(lambda: [0,0])

    # compter la poule A
    for joueur in pouleA:
        counts[joueur['club']][0] += 1

    # compter la poule B
    for joueur in pouleB:
        counts[joueur['club']][1] += 1

    return dict(counts)

def echanger_joueur(poule_surplus, poule_autre, club):
    """
    Cherche un joueur de 'club' dans la poule_surplus (à l'indice != 0),
    et un joueur de l'autre poule (indice != 0) d'un club différent,
    pour tenter l'échange.
    Retourne True si échange réussi, False sinon.
    """

    # Indices sur lesquels on a le droit d'agir (1, 2, 3) seulement
    indices_surplus = range(1, len(poule_surplus))
    indices_autre = range(1, len(poule_autre))

    # On va essayer toutes les paires (Jsurplus, Jautre)
    for i in indices_surplus:
        j1 = poule_surplus[i]
        if j1['club'] != club:
            continue  # on cherche un joueur QUI est de ce club pour le déplacer
        
        # On parcourt les joueurs de l'autre poule
        for k in indices_autre:
            j2 = poule_autre[k]
            if j2['club'] == club:
                continue  # on veut échanger contre un joueur d'un club différent
            
            # calculer la représentation du club j dans les 2 poules
            club_j2 = j2['club']
            count_j2_in_pouleA = sum(1 for j in poule_surplus if j['club'] == club_j2) + 1
            count_j2_in_pouleB = sum(1 for j in poule_autre if j['club'] == club_j2) - 1
            
            if abs(count_j2_in_pouleA - count_j2_in_pouleB) <= 1:
                # TENTATIVE d'échange
                print(f"on échange {j1} avec {j2}")
                poule_surplus[i], poule_autre[k] = j2, j1  # swap
                return True
            # Vérifier si ça a amélioré la situation ou si ça ne crée pas 
            # un nouveau déséquilibre sur un autre club 
            # => On peut recalculer le difference sur ce "club" tout de suite.
            #    Ou (plus complet) on peut vérifier tout le dict. 
            # Pour un code succinct, on suppose que l'échange est OK si on a 
            # effectivement déplacé un joueur du club "club".

            # On peut s'arrêter ici et laisser la boucle while 
            # de reequilibrer_2_poules re-tester pour d'autres clubs potentiels
            #return True

    # Si on n'a pas trouvé d'échange possible
    return False

def essayer_echange(pouleA, pouleB, club):
    """
    Tente de résoudre un surplus de 'club' dans pouleA ou pouleB 
    via un échange de joueurs. 
    Retourne True si un échange a été effectué, False sinon.
    """

    # 1) Qui a le surplus ?
    # On compte le club dans chaque poule
    countA = sum(1 for j in pouleA if j['club'] == club)
    countB = sum(1 for j in pouleB if j['club'] == club)

    # Cas: surplus en pouleA
    if countA > countB:
        return echanger_joueur(pouleA, pouleB, club)
    # Cas: surplus en pouleB
    elif countB > countA:
        return echanger_joueur(pouleB, pouleA, club)
    else:
        # pas de surplus en fait...
        return False
    
def reequilibrer_2_poules(pouleA, pouleB):
    """
    Tentative de rééquilibrage des clubs entre deux poules de 4 joueurs.
    Ne déplace jamais l'indice 0 (le seed #1) dans chaque poule.
    Modifie potentiellement pouleA, pouleB sur place.
    """
    # On itère tant qu'il y a un rééquilibrage possible
    # (ou qu'on fait un tour pour voir s'il reste un écart > 1 impossible à corriger).
    
    while True:
        # 1) On construit un dict des effectifs pour chaque club
        club_counts = construire_club_counts(pouleA, pouleB)
        print("club_counts: ", club_counts)
        for idx in pouleA:
            print(f"pouleA, club: {idx['club']}, dossard: {idx['dossard']}")
        for idx in pouleB:
            print(f"pouleB, club: {idx['club']}, dossard: {idx['dossard']}")
        
        # 2) On cherche s'il existe un club dont l'écart de counts > 1
        club_trouve = None
        for club, (countA, countB) in club_counts.items():
            if abs(countA - countB) > 1:
                club_trouve = club
                break
        
        if club_trouve is None:
            # Plus aucun club en déséquilibre > 1 => on arrête
            break
        
        # 3) On essaie de faire un échange pour ce club
        print(f"club {club_trouve} : surplus dans pouleA ? {club_counts[club_trouve][0]} > {club_counts[club_trouve][1]}")
        reussi = essayer_echange(pouleA, pouleB, club_trouve)
        print(f"club {club_trouve} : échange réussi ? {reussi}")
        if not reussi:
            # Aucun échange n'est possible, on arrête
            break

    # A la fin, pouleA et pouleB sont potentiellement rééquilibrées
    return pouleA, pouleB


def construire_tops_2_poules(joueurs_top):
    top8 = joueurs_top.iloc[:8]
    indices_top1 = [0, 3, 4, 7]
    indices_top2 = [1, 2, 5, 6]

    df_poule_top1 = top8.iloc[indices_top1]
    df_poule_top2 = top8.iloc[indices_top2]

    poule_top1 = df_poule_top1.to_dict('records')
    poule_top2 = df_poule_top2.to_dict('records')
    print("poule_top1: ", poule_top1)
    print("poule_top2: ", poule_top2)

    # rééquilibrage
    poule_top1, poule_top2 = reequilibrer_2_poules(poule_top1, poule_top2)
    print("poule_top1 après rééquilibrage: ", poule_top1)
    print("poule_top2 après rééquilibrage: ", poule_top2)

    return poule_top1, poule_top2

def reequilibrer_apres_echec_echange(poules, club_limits, club_min):
    """
    Rééquilibre les poules après un échec d'échange en respectant les contraintes :
    - Chaque club doit respecter son minimum (`club_min`) et son maximum (`club_limits`) dans chaque poule.
    - Le nombre de poules reste constant.
    """
    nb_poules = len(poules)

    def count_in_poule(poule, club):
        return sum(1 for j in poule if j["club"] == club)

    changed = True
    while changed:
        changed = False

        # Parcourir toutes les poules
        for idx_x, pouleX in enumerate(poules):
            # On repère s'il y a un surplus ou un manque pour un club C
            from collections import Counter
            counts = Counter(j["club"] for j in pouleX)
            surplus_clubs = [(c, n) for (c, n) in counts.items() if n > club_limits.get(c, 3)]
            missing_clubs = [(c, n) for c, n in club_min.items() if counts.get(c, 0) < n]

            #print(f"POULE {idx_x} : surplus_clubs: {surplus_clubs}, missing_clubs: {missing_clubs}")
            # Traiter les clubs en surplus
            for (clubC, countC) in surplus_clubs:
                maxC = club_limits.get(clubC, 3)

                # Descendre la pouleX du bas vers le haut
                found_exchange = False
                for posC in range(len(pouleX)-1, -1, -1):
                    if pouleX[posC]["club"] == clubC:
                        joueurC = pouleX[posC]

                        # Chercher une poule Y où count_in_poule(Y, clubC) < maxC
                        for idx_y, pouleY in enumerate(poules):
                            if idx_y == idx_x:
                                continue
                            countC_in_y = count_in_poule(pouleY, clubC)
                            if countC_in_y >= maxC:
                                continue  # pas de place pour clubC ici

                            # On parcourt pouleY du bas vers le haut
                            for posD in range(len(pouleY)-1, -1, -1):
                                joueurD = pouleY[posD]
                                clubD = joueurD["club"]
                                maxD = club_limits.get(clubD, 3)
                                minD = club_min.get(clubD, 0)

                                # Vérif qu'en envoyant joueurC => pouleY,
                                # on ne dépasse pas la limite pour clubC dans Y
                                nouveau_countC_in_y = countC_in_y
                                if clubC != clubD:
                                    nouveau_countC_in_y += 1
                                if nouveau_countC_in_y > maxC:
                                    continue

                                # Vérif qu'en envoyant joueurD => pouleX,
                                # on ne dépasse pas la limite pour clubD dans X
                                countD_in_x = count_in_poule(pouleX, clubD)
                                nouveau_countD_in_x = countD_in_x
                                if clubD != clubC:
                                    nouveau_countD_in_x += 1
                                if nouveau_countD_in_x > maxD:
                                    continue

                                # Vérif qu'on respecte les minimums pour clubC et clubD
                                if countC - 1 < club_min.get(clubC, 0) or countD_in_x + 1 < minD:
                                    continue

                                # Échange possible
                                pouleX[posC], pouleY[posD] = joueurD, joueurC
                                print(f"[reequilibrer_apres_echec_echange] ÉCHANGE :"
                                      f" Poule {idx_x} pos {posC} (club {clubC}) <--> "
                                      f" Poule {idx_y} pos {posD} (club {clubD})")
                                print(f"    JoueurC : {joueurC['nom']} {joueurC['prenom']} (club {clubC})")
                                print(f"    JoueurD : {joueurD['nom']} {joueurD['prenom']} (club {clubD})")

                                changed = True
                                found_exchange = True
                                break  # on arrête la boucle pouleY[posD]
                            if found_exchange:
                                break  # on arrête la boucle pouleY
                        if found_exchange:
                            break  # on arrête la boucle posC

            # Traiter les clubs en manque
            for (clubC, minC) in missing_clubs:
                countC = counts.get(clubC, 0)

                # Chercher une poule Y où count_in_poule(Y, clubC) > minC
                for idx_y, pouleY in enumerate(poules):
                    print (f"idx_y: {idx_y}")
                    if idx_y == idx_x:
                        continue
                    countC_in_y = count_in_poule(pouleY, clubC)
                    if countC_in_y <= minC:
                        continue  # pas assez de joueurs pour clubC ici
  
                    # On parcourt pouleY du bas vers le haut
                    for posD in range(len(pouleY)-1, -1, -1):
                        joueurD = pouleY[posD]
                        clubD = joueurD["club"]
                        maxD = club_limits.get(clubD, 3)
                        minD = club_min.get(clubD, 0)
                        if (posD+1 > len(pouleX)):
                            posC = len(pouleX) - 1
                        else:
                            posC = posD
                        joueurC = pouleX[posC]
                        clubCurC = joueurC["club"]

                        # Vérif qu'en envoyant joueurD => pouleX,
                        # on respecte le minimum pour clubC dans X
                        if countC + 1 < minC:
                            continue

                        # Vérif qu'en envoyant joueurD => pouleX,
                        # on ne dépasse pas la limite pour clubD dans X
                        countD_in_x = count_in_poule(pouleX, clubD)
                        if countD_in_x + 1 > maxD:
                            continue
                        
                        print(f"clubC: {clubC}, clubD: {clubD}")
                        
                        # Échange possible
                        #pouleX.append(joueurD)
                        #pouleY.pop(posD)
                        
                        if (clubC != clubD):
                            continue
                        
                        pouleX[posC], pouleY[posD] = joueurD, joueurC
                        print(f"[reequilibrer_apres_echec_echange] ECHANGE :"
                              f" Poule {idx_x} reçoit {joueurD['nom']} {joueurD['prenom']} (club {clubD})")

                        changed = True
                        break  # on arrête la boucle pouleY[posD]
                    if changed:
                        break  # on arrête la boucle pouleY
    return poules

def verifier_repartition_clubs(poules, club_limits, club_min):
    """
    Vérifie si, dans chaque poule, aucun club ne dépasse sa limite 'club_limits[club]'.
    Retourne True s'il reste un problème (un surplus), False si tout est correct.
    Affiche en console un message si on détecte un surplus.
    """

    from collections import Counter
    liste_clubs = []
    for idx_poule, poule in enumerate(poules):
        for joueur in poule:
            if joueur["club"] not in liste_clubs:
                liste_clubs.append(joueur["club"])
    
    probleme = False

    for idx_poule, poule in enumerate(poules):
        # Construire un dict {club: count} pour cette poule
        counts = Counter(j["club"] for j in poule)

        # Vérifier chaque club
        for club, nb_joueurs in counts.items():
            limite = club_limits.get(club, 3)  # par défaut 3
            min = club_min.get(club, 3)
            if nb_joueurs > limite:
                print(f"[verifier_repartition_clubs] Surplus dans poule {idx_poule}: "
                      f"Club {club} a {nb_joueurs} joueurs (limite = {limite}).")
                probleme = True
            if nb_joueurs < min:
                print(f"[verifier_repartition_clubs] Manque dans poule {idx_poule}: "
                      f"Club {club} a {nb_joueurs} joueurs (min = {min}).")
                probleme = True
                
        for club,min in club_min.items():
            if min>0 and club not in counts:
                print(f"[verifier_repartition_clubs] Manque dans poule {idx_poule}: "
                      f"Club {club} a 0 joueurs (min = {min}).")
                probleme = True
    return probleme

'''
def reequilibrer_poules_par_echange(poules, club_limits, start_index=1):
    """
    Parcourt chaque poule X. Pour chaque club C qui dépasse sa limite,
    on parcourt la poule X en cherchant un joueur de club C (à partir de la position start_index).
    Puis on tente un échange avec la poule Y, à la même position (posC).
    On vérifie que déplacer le joueur de club C vers Y ne dépasse pas la limite de C dans Y,
    et que déplacer le joueur de Y (club D) vers X ne dépasse pas la limite de D dans X.
    Si ça résout le surplus, on arrête pour ce club. Sinon, on continue à chercher un autre joueur C.
    
    Les prints donnent des informations sur les échanges effectués (joueurs, poules, positions).
    """

    nb_poules = len(poules)

    def count_in_poule(poule, club):
        return sum(1 for j in poule if j["club"] == club)

    changed = True
    while changed:
        changed = False

        # Parcourir chaque poule X
        for x in range(nb_poules):
            pouleX = poules[x]

            # Construire un "dictionnaire" club->count pour la poule X
            from collections import Counter
            club_counts = Counter(j["club"] for j in pouleX)

            # Pour chaque club C présent dans la poule
            for clubC, countC in club_counts.items():
                maxC = club_limits.get(clubC, 3)  # par défaut 3 si non défini

                # Tant que ce club dépasse la limite
                while countC > maxC:
                    # On parcourt la poule X pour trouver un joueur de clubC
                    # à partir de la position start_index
                    posC = -1
                    for idxC in range(start_index, len(pouleX)):
                        if pouleX[idxC]["club"] == clubC:
                            posC = idxC
                            break

                    if posC == -1:
                        # On n'a pas trouvé d'autre joueur "déplaçable" pour ce club
                        break  # on passe au club suivant

                    joueurC = pouleX[posC]
                    clubC = joueurC["club"]

                    # Tenter un échange avec une poule Y
                    echange_trouve = False
                    for y in range(nb_poules):
                        if y == x:
                            continue  # pas d'échange dans la même poule
                        pouleY = poules[y]

                        # Vérifier qu'il existe bien la même position posC dans pouleY
                        if len(pouleY) <= posC:
                            continue  # pas de joueur à cet index dans pouleY

                        joueurD = pouleY[posC]
                        clubD = joueurD["club"]
                        maxD = club_limits.get(clubD, 3)

                        # Vérif 1 : pouleY ne dépasse pas la limite pour clubC
                        countC_in_y = count_in_poule(pouleY, clubC)

                        # Cas clubC != clubD => on ferait +1 au clubC
                        if clubC != clubD:
                            nouveau_countC_in_y = countC_in_y + 1
                        else:
                            # Même club => échanger un joueurC contre un joueurD du même club
                            # ne change pas la quantité de ce club dans Y
                            nouveau_countC_in_y = countC_in_y

                        if nouveau_countC_in_y > maxC:
                            # Dépasse la limite => pas possible
                            continue

                        # Vérif 2 : pouleX ne dépasse pas la limite pour clubD
                        countD_in_x = count_in_poule(pouleX, clubD)
                        if clubD != clubC:
                            nouveau_countD_in_x = countD_in_x + 1
                        else:
                            # Même club => l'échange ne modifie pas la distribution de ce club
                            # dans X, on demeure en surplus pour clubC ?
                            # Selon votre logique, on skip.
                            continue

                        if nouveau_countD_in_x > maxD:
                            # On dépasserait la limite pour clubD => pas possible
                            continue

                        # Si on arrive ici, l'échange est possible
                        pouleX[posC], pouleY[posC] = joueurD, joueurC
                        print(f"[reequilibrer_poules_par_echange] ÉCHANGE :"
                              f" Poule {x} pos {posC} (club {clubC}) <-->"
                              f" Poule {y} pos {posC} (club {clubD})")
                        print(f"    JoueurC : {joueurC['nom']} {joueurC['prenom']} (club {clubC})")
                        print(f"    JoueurD : {joueurD['nom']} {joueurD['prenom']} (club {clubD})")

                        changed = True
                        echange_trouve = True

                        # recalculer le count du clubC dans pouleX
                        countC = count_in_poule(pouleX, clubC)
                        break  # on arrête la boucle sur pouleY

                    if not echange_trouve:
                        # On n'a pas trouvé de poule Y pour échanger
                        break
    return poules

'''
def reequilibrer_poules_par_echange(poules, club_limits, club_min, start_index=1):
    """
    Parcourt chaque poule X. Pour chaque club C qui dépasse sa limite ou ne respecte pas son minimum,
    on parcourt la poule X en cherchant un joueur de club C (à partir de la position start_index).
    Puis on tente un échange avec la poule Y, à la même position (posC).
    On vérifie que déplacer le joueur de club C vers Y ne dépasse pas la limite de C dans Y,
    et que déplacer le joueur de Y (club D) vers X ne dépasse pas la limite de D dans X,
    tout en respectant les minimums définis par club_min.

    Les prints donnent des informations sur les échanges effectués (joueurs, poules, positions).
    """

    nb_poules = len(poules)

    def count_in_poule(poule, club):
        return sum(1 for j in poule if j["club"] == club)

    changed = True
    while changed:
        changed = False

        # Parcourir chaque poule X
        for x in range(nb_poules):
            pouleX = poules[x]

            # Construire un "dictionnaire" club->count pour la poule X
            from collections import Counter
            club_counts = Counter(j["club"] for j in pouleX)

            # Pour chaque club C présent dans la poule
            for clubC, countC in club_counts.items():
                maxC = club_limits.get(clubC, 3)  # par défaut 3 si non défini
                minC = club_min.get(clubC, 0)  # par défaut 0 si non défini

                # Tant que ce club dépasse la limite ou ne respecte pas le minimum
                while countC > maxC or countC < minC:
                    # On parcourt la poule X pour trouver un joueur de clubC
                    # à partir de la position start_index
                    posC = -1
                    for idxC in range(start_index, len(pouleX)):
                        if pouleX[idxC]["club"] == clubC:
                            posC = idxC
                            break

                    if posC == -1:
                        # On n'a pas trouvé d'autre joueur "déplaçable" pour ce club
                        break  # on passe au club suivant

                    joueurC = pouleX[posC]
                    clubC = joueurC["club"]

                    # Tenter un échange avec une poule Y
                    echange_trouve = False
                    for y in range(nb_poules):
                        if y == x:
                            continue  # pas d'échange dans la même poule
                        pouleY = poules[y]

                        # Vérifier qu'il existe bien la même position posC dans pouleY
                        if len(pouleY) <= posC:
                            continue  # pas de joueur à cet index dans pouleY

                        joueurD = pouleY[posC]
                        clubD = joueurD["club"]
                        maxD = club_limits.get(clubD, 3)
                        minD = club_min.get(clubD, 0)

                        # Vérif 1 : pouleY ne dépasse pas la limite pour clubC
                        countC_in_y = count_in_poule(pouleY, clubC)

                        # Cas clubC != clubD => on ferait +1 au clubC
                        if clubC != clubD:
                            nouveau_countC_in_y = countC_in_y + 1
                        else:
                            # Même club => échanger un joueurC contre un joueurD du même club
                            # ne change pas la quantité de ce club dans Y
                            nouveau_countC_in_y = countC_in_y

                        if nouveau_countC_in_y > maxC:
                            # Dépasse la limite => pas possible
                            continue

                        # Vérif 2 : pouleX ne dépasse pas la limite pour clubD
                        countD_in_x = count_in_poule(pouleX, clubD)
                        if clubD != clubC:
                            nouveau_countD_in_x = countD_in_x + 1
                        else:
                            # Même club => l'échange ne modifie pas la distribution de ce club
                            # dans X, on demeure en surplus pour clubC ? Selon votre logique, on skip.
                            continue

                        if nouveau_countD_in_x > maxD:
                            # On dépasserait la limite pour clubD => pas possible
                            continue

                        # Vérif 3 : Respect des minimums pour clubC et clubD
                        if countC - 1 < minC or countD_in_x + 1 < minD:
                            # L'échange créerait un manque pour l'un des clubs => pas possible
                            continue

                        # Si on arrive ici, l'échange est possible
                        pouleX[posC], pouleY[posC] = joueurD, joueurC
                        print(f"[reequilibrer_poules_par_echange] ÉCHANGE :"
                              f" Poule {x} pos {posC} (club {clubC}) <-->"
                              f" Poule {y} pos {posC} (club {clubD})")
                        print(f"    JoueurC : {joueurC['nom']} {joueurC['prenom']} (club {clubC})")
                        print(f"    JoueurD : {joueurD['nom']} {joueurD['prenom']} (club {clubD})")

                        changed = True
                        echange_trouve = True

                        # recalculer le count du clubC dans pouleX
                        countC = count_in_poule(pouleX, clubC)
                        break  # on arrête la boucle sur pouleY

                    if not echange_trouve:
                        # On n'a pas trouvé de poule Y pour échanger
                        break
    return poules

def placer_joueurs(autres_joueurs, nb_poules_3, nb_poules_4, nb_poules_5):
    """
    Crée des poules (3,4,5 joueurs) et y place tous les joueurs
    selon un 'ordre serpent' principal, mais s'adapte si la poule indiquée est pleine
    en essayant la suivante, etc., jusqu'à trouver un slot libre.

    - autres_joueurs : dataframe ou liste de dict (déjà trié).
    - nb_poules_3, nb_poules_4, nb_poules_5 : nb de poules de chaque taille.

    Retourne un tableau 'poules' (liste de listes de dict).
    """

    # Convertir en liste si c'est un DataFrame
    listed_joueurs = (
        autres_joueurs.to_dict("records") 
        if hasattr(autres_joueurs, "to_dict") 
        else list(autres_joueurs)
    )

    # 1) Créer la liste des poules vides et un tableau 'capacites'
    poules = []
    capacites = []
    for _ in range(nb_poules_5):
        poules.append([])
        capacites.append(5)
    for _ in range(nb_poules_4):
        poules.append([])
        capacites.append(4)
    for _ in range(nb_poules_3):
        poules.append([])
        capacites.append(3)

    nb_poules = len(poules)
    nb_joueurs = len(listed_joueurs)
    
    club_counts = {}
    for j in autres_joueurs.to_dict("records"):
        c = j["club"]
        club_counts[c] = club_counts.get(c, 0) + 1

    club_limits = {}
    club_min = {}
    for (club, y) in club_counts.items():
        if y <= nb_poules:
            club_limits[club] = 1
            club_min[club] = 0
        elif y <= 2*nb_poules:
            club_limits[club] = 2
            club_min[club] = 1
        else:
            # Par défaut 3, voire plus si vous voulez
            club_limits[club] = 3
            club_min[club] = 2

    # 2) Construire l'ordre serpent
    serpent_order = construire_serpent_order(nb_poules, nb_joueurs)
    
    # 3) Parcourir chaque joueur
    for i, joueur in enumerate(listed_joueurs):
        p = serpent_order[i]  # poule "préférée" par l'ordre serpent

        # Tenter de la placer dans la poule p ; si c'est plein,
        # on avance p+1, etc. (modulo nb_poules) jusqu'à trouver un slot.
        start_p = p
        placed = False
        while True:
            if len(poules[p]) < capacites[p]:
                # Il y a de la place
                poules[p].append(joueur)
                placed = True
                break
            else:
                # Avancer la poule (en mode circulaire)
                p = (p + 1) % nb_poules
                # Si on est revenu au point de départ, plus aucun slot libre => saturé
                if p == start_p:
                    break
        
        if not placed:
            print(f"Attention : plus aucun slot libre pour {joueur['nom']} {joueur['prenom']}")
            # On pourrait lever une Exception ou faire un break, 
            # selon vos besoins
            # raise ValueError("Impossible de placer tous les joueurs - pas assez de place.")
    reequilibrer_poules_par_echange(poules, club_limits, club_min, start_index=1)
    reequilibrer_apres_echec_echange(poules, club_limits, club_min)
    surplus = verifier_repartition_clubs(poules, club_limits, club_min)
    if surplus:
        print("==> Il reste un surplus après tentatives d'échanges.")
    else:
        print("==> Tout est OK côté répartition de clubs.")

    return poules


def construire_serpent_order(nb_poules, nb_joueurs):
    """
    Construit un tableau d'indices de poule pour un 'serpent' sur nb_joueurs.
    Ex.: si nb_poules=3 et nb_joueurs=8 => [0,1,2,1,0,1,2,1].
    """
    order = []
    idx = 0
    direction = 1
    for _ in range(nb_joueurs):
        order.append(idx)
        new_idx = idx + direction
        if new_idx >= nb_poules:
            new_idx = nb_poules - 1
            direction = -1
        elif new_idx < 0:
            new_idx = 0
            direction = 1
        idx = new_idx
    return order


def ajuster_espace_autour_tableau(table):
    # Paragraphe avant le tableau
    prev_paragraph = table._element.getprevious()
    if prev_paragraph is not None and prev_paragraph.tag.endswith("p"):
        prev_p = table._parent.paragraphs[-1]
        prev_p.paragraph_format.space_after = Pt(0)
        prev_p.paragraph_format.space_before = Pt(0)

    # Paragraphe après le tableau
    next_paragraph = table._element.getnext()
    if next_paragraph is not None and next_paragraph.tag.endswith("p"):
        next_p = table._parent.paragraphs[-1]
        next_p.paragraph_format.space_after = Pt(0)
        next_p.paragraph_format.space_before = Pt(0)

            
def ecrire_poule_on_frame(doc, titre, joueurs, is_top=False):
    oppositions = [(1, 4), (2, 3), (1, 3), (2, 4), (1, 2), (3, 4), (1, 5), (2,5)]
    heading = titre if is_top else f"Poule {titre}"

    paragraph = doc.add_paragraph()
    paragraph.text = f'{titre}'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Style du titre
    run = paragraph.runs[0]
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.bold = True

    #table = doc.add_table(rows=1, cols=5)
    table = doc.add_table(rows=1, cols=5)
    ajuster_espace_autour_tableau(table)
    table.autofit = False
    table.style = 'Table Grid'

    table.columns[0].width = Cm(0.4)
    table.columns[1].width = Cm(3.3)
    table.columns[2].width = Cm(3.3)
    table.columns[3].width = Cm(1.5)
    table.columns[4].width = Cm(1)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = ' '
    hdr_cells[1].text = 'Nom'
    hdr_cells[2].text = 'Prenom'
    hdr_cells[3].text = 'Club'
    hdr_cells[4].text = 'Dos.'

    for idx, joueur in enumerate(joueurs, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = joueur['nom'] if isinstance(joueur, dict) else joueur.nom
        row_cells[2].text = joueur['prenom'] if isinstance(joueur, dict) else joueur.prenom
        row_cells[3].text = joueur['club'] if isinstance(joueur, dict) else joueur.club
        row_cells[4].text = str(int(joueur['dossard'])) if isinstance(joueur, dict) else str(int(joueur.dossard))

    table2 = doc.add_table(rows=1, cols=6 + len(joueurs))
    table2.style = 'Table Grid'
    ajuster_espace_autour_tableau(table2)
    table2.columns[0].width = Cm(1)
   
    for i in range(1, 6 + len(joueurs)):
            if i < len(table2.columns):
                table2.columns[i].width = Cm(0.8)

    hdr_cells = table2.rows[0].cells
    hdr_cells[6].text = "1"
    hdr_cells[7].text = "2"
    hdr_cells[8].text = "3"
    if len(joueurs) == 4:
            hdr_cells[9].text = "4"

    for z in range(len(oppositions)):
            x, y = oppositions[z]
            if x <= len(joueurs) and y <= len(joueurs):
                row_cells = table2.add_row().cells
                rencontre = f'{x} vs {y}'
                row_cells[0].text = rencontre
            for k in range (1,len(joueurs) + 1):
                if (k != x and k != y and y<=len(joueurs)):
                    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
                    row_cells[k+5]._tc.get_or_add_tcPr().append(shading_elm_1)
    row_cells = table2.add_row().cells
    for z in range (0, 5):
        row_cells[z].text = ""
    row_cells[5].text = "Tot."
    
    # Suppression des bordures pour les cellules avant "Tot."
    for i in range(5):  # Les 5 premières colonnes avant "Tot."
        tc = row_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()

        # Création de l'élément XML pour supprimer les bordures
        borders = OxmlElement('w:tcBorders')

        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')  # Utilisation correcte de qn()
            borders.append(border)

        tcPr.append(borders)
    
    table3 = doc.add_table(rows=1, cols=2)
    ajuster_espace_autour_tableau(table3)
    table3.style = 'Table Grid'
    table3.columns[0].width = Cm(1)
    table3.columns[1].width = Cm(8)

    row_cells = table3.rows[0].cells
    for index in range(1, len(joueurs) + 1):
        if index == 1:
          label = "1er"
        else:
           label = f"{index}eme"

        row_cells[0].text = f'{label}'

        # Appliquer une hauteur de ligne personnalisée
        tr = row_cells[0]._tc.getparent()  # Obtenir l'élément XML de la ligne
        trPr = OxmlElement('w:trPr')  # Créer une propriété de ligne
        trHeight = OxmlElement('w:trHeight')  # Définir la hauteur de ligne
        trHeight.set(qn('w:val'), "500")  # Augmenter la hauteur (valeur en 1/20 de point)
        trPr.append(trHeight)
        tr.append(trPr)  # Appliquer la propriété à la ligne

        if index < len(joueurs):
            row_cells = table3.add_row().cells
        
def create_word_file_with_frames(doc, groupes, categorie, first_poule, is_top, position=0, table=None):
    # Configurer les marges de la page
    section = doc.sections[0]
    section.page_width = Inches(8.27)  # Largeur de la page (A4)
    section.page_height = Inches(11.69)  # Hauteur de la page (A4)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    # Si aucune table n'existe, en créer une nouvelle
    if table is None:
        table = doc.add_table(rows=3, cols=2)
        table.autofit = False  # Désactiver l'ajustement automatique
        for column in table.columns:
            column.width = Inches(3.75)

    for i, joueurs in enumerate(groupes, start=1):
        # Ajouter une nouvelle page après avoir rempli la table
        if position == 0 and i > 0:
            if not first_poule:
                doc.add_page_break()
                table = doc.add_table(rows=3, cols=2)
                table.autofit = False
                for column in table.columns:
                    column.width = Inches(3.75)
            else:
                first_poule = False

        # Déterminer la cellule où écrire
        if position == 0:  # Haut gauche
            cell = table.cell(0, 0)
        elif position == 1:  # Haut droite
            cell = table.cell(0, 1)
        elif position == 2:  # Milieu gauche
            cell = table.cell(1, 0)
        elif position == 3:  # Milieu droite
            cell = table.cell(1, 1)
        #elif position == 4:  # Bas gauche
        #    cell = table.cell(2, 0)
        #elif position == 5:  # Bas droite
        #    cell = table.cell(2, 1)
       
        if is_top:
            titre=f"Top {i}"
        else:
            titre=f"Poule {i}"
        titrePoule = f"{titre} {categorie}"
        ecrire_poule_on_frame(cell, titrePoule, joueurs, is_top)

        # Mettre à jour la position
        #position = (position + 1) % 6  # Cycle entre 0 et 5
        position = (position + 1) % 4  # Cycle entre 0 et 5

    return position, first_poule, table  # Retourner la position actuelle et la table en cours

def ecrire_document_word(fichier_sortie, feuille_source_resultat, top_poules, poules):
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = "Arial"
    font.size = Pt(7)

    first_poule = True
    
    # Initialiser la position à 0 pour démarrer avec les poules de top
    position, first_poule, table = create_word_file_with_frames(doc, top_poules, feuille_source_resultat, first_poule, is_top=True)

    # Écrire les poules normales à partir de la dernière position et table
    position, _, _ = create_word_file_with_frames(doc, poules, feuille_source_resultat, first_poule, is_top=False, position=position, table=table)

    fichier_sortie = lire_word_name(fichier_sortie)
    doc.save(fichier_sortie)
    print(f'Fichier Word généré: {fichier_sortie}')

def repartir_poules_sans_conflit(fichier_entree, feuille_source_resultat, nb_poules_top, nb_poules_3=None, nb_poules_4=None, nb_poules_5=None):
    joueurs = charger_joueurs(fichier_entree, feuille_source_resultat)
    joueurs, tourActif, nbJoueurs = filtrer_et_trier_joueurs(joueurs, nb_poules_top)

    print("tourActif:", tourActif, ", nbJoueurs:", nbJoueurs)

    if (nb_poules_3 != None and nb_poules_4 != None):
        nbTotal = nb_poules_top*4 + nb_poules_3*3 + nb_poules_4*4
        if (nb_poules_5 != None):
            nbTotal += nb_poules_5 * 5
        if (nbTotal != nbJoueurs):
            print(nbJoueurs, " joueurs inscrits, erreur dans la definition des parametres (nb_poules_top: ", nb_poules_top, ", nb_poules_3: ", nb_poules_3, ", nb_poules_4", nb_poules_4)
            exit()
    if nb_poules_top > 0:
        top_poules, autres_joueurs = construire_tops(joueurs, nb_poules_top)
        print("after")
    else:
        top_poules = {}
        autres_joueurs = joueurs

    if nb_poules_3 is None or nb_poules_4 is None:
        total_joueurs = int(len(autres_joueurs))
        nb_poules_4 = total_joueurs // 4
        nb_poules_3 = (total_joueurs % 4) // 3
        
        while ( (nb_poules_4*4 + nb_poules_3*3) < total_joueurs):
            nb_poules_4 = nb_poules_4 - 1
            nb_poules_3 = int( (total_joueurs - (nb_poules_4 *4))/3)
        
    print ("Nb Poules de 3", nb_poules_3, ", Nb Poules de 4", nb_poules_4, ", Nb Poules de 5", nb_poules_5)

    poules = placer_joueurs(autres_joueurs, nb_poules_3, nb_poules_4, nb_poules_5)
    '''fichierWord=f'{feuille_source_resultat}_{tourActif}.docx'
    ecrire_document_word(fichierWord, feuille_source_resultat, top_poules, poules)
    # Générer le fichier JSON
    ecrire_fichier_json(feuille_source_resultat, tourActif, top_poules, poules)
    '''
    return top_poules, poules, tourActif, nbJoueurs

def update_word_and_json(categorie, tourActif, top_poules, poules):

    fichierWord=f'{categorie}_{tourActif}.docx'
    ecrire_document_word(fichierWord, categorie, top_poules, poules)
    # Générer le fichier JSON
    ecrire_fichier_json(categorie, tourActif, top_poules, poules)
    
    messagebox.showinfo("Information", "Le fichier Word et le fichier JSON ont été mis à jour avec succès.")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Repartir les poules en fonction des joueurs et des options.")
    parser.add_argument("Categorie", type=str, help="Nom de la feuille source dans le fichier Excel.")
    parser.add_argument("--nb_poules_top", type=int, default=None, help="Nombre de poules Top. Optionnel.")
    parser.add_argument("--nb_poules_3", type=int, default=None, help="Nombre de poules de 3 joueurs. Optionnel.")
    parser.add_argument("--nb_poules_4", type=int, default=None, help="Nombre de poules de 4 joueurs. Optionnel.")
    parser.add_argument("--nb_poules_5", type=int, default=None, help="Nombre de poules de 5 joueurs. Optionnel.")
    args = parser.parse_args()
    
    fichier_entree = lire_excel_name()
    print(f"Fichier complet : {fichier_entree}")
        
    nb_tops_defaut = int(charger_config_categorie("config.ini", args.Categorie))
        
    number_of_tops=args.nb_poules_top
    if args.nb_poules_top == None:
        number_of_tops = nb_tops_defaut
    
    total_poules_5 = 0
    if args.nb_poules_5 != None:
        total_poules_5 = args.nb_poules_5
   
    print(f"number_of_tops: {number_of_tops}")
   
    repartir_poules_sans_conflit(
        fichier_entree,
        feuille_source_resultat=args.Categorie,
        nb_poules_top=number_of_tops,
        nb_poules_3=args.nb_poules_3,
        nb_poules_4=args.nb_poules_4,
        nb_poules_5=total_poules_5)
